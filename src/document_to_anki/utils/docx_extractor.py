"""DOCX text extraction utilities."""

from pathlib import Path

from docx import Document
from loguru import logger

from .text_extractor_common import TextExtractionError


def extract_text(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = Document(str(file_path))
        text_content: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        extracted_text = "\n\n".join(text_content)
        if not extracted_text.strip():
            logger.warning(f"No text content extracted from DOCX: {file_path}")
            return ""

        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return extracted_text
    except FileNotFoundError as e:
        raise TextExtractionError(f"DOCX file not found: {file_path}") from e
    except PermissionError as e:
        raise TextExtractionError(f"Permission denied accessing DOCX file: {file_path}") from e
    except Exception as e:
        error_msg = str(e).lower()
        if "not a zip file" in error_msg or "bad zipfile" in error_msg:
            raise TextExtractionError(f"Invalid or corrupted DOCX file: {file_path}") from e
        if "no such file" in error_msg:
            raise TextExtractionError(f"DOCX file not found: {file_path}") from e
        raise TextExtractionError(f"Unexpected error extracting from DOCX {file_path}: {str(e)}") from e
